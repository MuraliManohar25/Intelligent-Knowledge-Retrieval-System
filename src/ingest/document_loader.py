"""Unified document loader supporting PDF and DOCX files."""
import PyPDF2
from docx import Document
from pathlib import Path
from typing import List, Dict
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoader:
    """Load and extract text from PDF and DOCX documents."""
    
    def __init__(self, documents_dir: Path):
        self.documents_dir = Path(documents_dir)
    
    def load_pdf(self, pdf_path: Path) -> List[Dict[str, any]]:
        """
        Load a single PDF and extract text page by page.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of dicts with page_num and text
        """
        pages = []
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                logger.info(f"Loading {pdf_path.name}: {num_pages} pages")
                
                for page_num in range(num_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        
                        if text.strip():  # Only add if text exists
                            pages.append({
                                'page_num': page_num + 1,  # 1-indexed
                                'text': text,
                                'source_file': pdf_path.name
                            })
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error loading PDF {pdf_path}: {e}")
            return []
        
        logger.info(f"Successfully loaded {len(pages)} pages from {pdf_path.name}")
        return pages
    
    def load_docx(self, docx_path: Path) -> List[Dict[str, any]]:
        """
        Load a single DOCX file and extract text.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            List of dicts with page_num and text (simulated pages)
        """
        pages = []
        
        try:
            doc = Document(docx_path)
            logger.info(f"Loading {docx_path.name}")
            
            # Combine all paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Combine all text
            document_text = "\n\n".join(full_text)
            
            if not document_text.strip():
                logger.warning(f"Document {docx_path.name} appears to be empty (no text content found)")
                return []
            
            # Split into logical "pages" (approximately 2000 characters per page)
            # This simulates page breaks for consistency with PDF handling
            chars_per_page = 2000
            text_length = len(document_text)
            num_pages = max(1, (text_length + chars_per_page - 1) // chars_per_page)
            
            for page_num in range(num_pages):
                start_idx = page_num * chars_per_page
                end_idx = min((page_num + 1) * chars_per_page, text_length)
                page_text = document_text[start_idx:end_idx]
                
                if page_text.strip():
                    pages.append({
                        'page_num': page_num + 1,
                        'text': page_text,
                        'source_file': docx_path.name
                    })
            
            if pages:
                logger.info(f"Successfully loaded {len(pages)} pages from {docx_path.name}")
            else:
                logger.warning(f"No pages extracted from {docx_path.name}")
            
        except Exception as e:
            logger.error(f"Error loading DOCX {docx_path}: {e}")
            return []
        
        return pages
    
    def load_document(self, file_path: Path) -> List[Dict[str, any]]:
        """
        Load a document (PDF or DOCX) based on file extension.
        
        Args:
            file_path: Path to document file
            
        Returns:
            List of dicts with page_num and text
        """
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self.load_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.load_docx(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext} for {file_path.name}")
            return []
    
    def load_all_documents(self) -> Dict[str, List[Dict]]:
        """
        Load all PDF and DOCX files from documents directory.
        
        Returns:
            Dict mapping filename to list of page dicts
        """
        all_documents = {}
        
        # Find all supported files
        pdf_files = list(self.documents_dir.glob("*.pdf"))
        docx_files = list(self.documents_dir.glob("*.docx"))
        doc_files = list(self.documents_dir.glob("*.doc"))
        
        all_files = pdf_files + docx_files + doc_files
        
        if not all_files:
            logger.warning(f"No supported documents found in {self.documents_dir}")
            logger.info("Supported formats: .pdf, .docx, .doc")
            return {}
        
        logger.info(f"Found {len(all_files)} document files")
        logger.info(f"  PDFs: {len(pdf_files)}, DOCX: {len(docx_files)}, DOC: {len(doc_files)}")
        
        for file_path in all_files:
            try:
                pages = self.load_document(file_path)
                if pages:
                    all_documents[file_path.name] = pages
                else:
                    logger.warning(f"Skipping {file_path.name} - no content extracted")
            except Exception as e:
                logger.error(f"Failed to load {file_path.name}: {e}")
                continue
        
        if not all_documents:
            logger.error("No documents were successfully loaded. Please check:")
            logger.error("  1. Files contain actual text content (not empty)")
            logger.error("  2. Files are valid PDF or DOCX format")
            logger.error("  3. Files are not corrupted")
        
        return all_documents

# Test function
if __name__ == "__main__":
    from config import DOCUMENTS_DIR
    loader = DocumentLoader(DOCUMENTS_DIR)
    docs = loader.load_all_documents()
    print(f"Loaded {len(docs)} documents")
    for doc_name, pages in docs.items():
        print(f"  {doc_name}: {len(pages)} pages")

